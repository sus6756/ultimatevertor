import streamlit as st
import tempfile
import pandas as pd
import subprocess
import os
from pdf2docx import Converter
from docx import Document
from PIL import Image

# ===== SAFE MOVIEPY IMPORT =====
try:
    from moviepy.editor import VideoFileClip, AudioFileClip
    VIDEO_ENABLED = True
except:
    VIDEO_ENABLED = False

st.set_page_config(page_title="Ultimate Converter", layout="wide")

# ===== UI =====
st.markdown("""
<style>
.big-title {font-size:42px;text-align:center;font-weight:bold;}
.sub-text {text-align:center;color:gray;margin-bottom:20px;}
</style>
""", unsafe_allow_html=True)

st.markdown('<div class="big-title">🔥 Ultimate Converter</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-text">Convert anything — fast & clean 🚀</div>', unsafe_allow_html=True)

# ===== TABS =====
tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "📄 Documents", "🖼️ Images", "📊 Data", "🎥 Video", "🎵 Audio"
])

# ===== HELPER =====
def save_temp_file(uploaded_file):
    ext = uploaded_file.name.split(".")[-1]
    temp = tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}")
    temp.write(uploaded_file.read())
    return temp.name, uploaded_file.name.split(".")[0]

# ================= DOCUMENTS =================
with tab1:
    st.subheader("📄 Document Converter")

    uploaded_file = st.file_uploader("Upload File", type=["pdf","docx","txt"])
    option = st.selectbox("Conversion", [
        "PDF to DOCX","DOCX to PDF","TXT to DOCX","DOCX to TXT"
    ])

    if uploaded_file:
        file_path, base_name = save_temp_file(uploaded_file)

        if st.button("🚀 Convert"):
            with st.spinner("Processing..."):
                try:
                    if option == "PDF to DOCX":
                        output = file_path.replace(".pdf", ".docx")
                        cv = Converter(file_path)
                        cv.convert(output)
                        cv.close()

                    elif option == "DOCX to PDF":
                        subprocess.run([
                            "soffice",
                            "--headless",
                            "--convert-to", "pdf",
                            file_path,
                            "--outdir", "/tmp"
                        ], check=True)

                        output = os.path.join(
                            "/tmp",
                            os.path.basename(file_path).replace(".docx", ".pdf")
                        )

                        if not os.path.exists(output):
                            st.error("❌ Conversion failed. File not created.")
                            st.stop()

                    elif option == "TXT to DOCX":
                        doc = Document()
                        with open(file_path,"r") as f:
                            doc.add_paragraph(f.read())
                        output = file_path + ".docx"
                        doc.save(output)

                    elif option == "DOCX to TXT":
                        doc = Document(file_path)
                        text = "\n".join([p.text for p in doc.paragraphs])
                        st.download_button("📥 Download TXT", text, file_name=f"{base_name}.txt")
                        st.stop()

                    with open(output,"rb") as f:
                        st.download_button(
                            "📥 Download",
                            f,
                            file_name=f"{base_name}_converted.{output.split('.')[-1]}"
                        )

                    st.success("✅ Done")

                except Exception as e:
                    st.error(f"❌ {e}")

# ================= IMAGE =================
with tab2:
    st.subheader("🖼️ Image → PDF")

    uploaded_file = st.file_uploader("Upload Image", type=["png","jpg","jpeg"])

    if uploaded_file:
        file_path, base_name = save_temp_file(uploaded_file)

        if st.button("🚀 Convert Image"):
            image = Image.open(file_path)
            output = file_path + ".pdf"
            image.convert("RGB").save(output)

            with open(output,"rb") as f:
                st.download_button("📥 Download PDF", f, file_name=f"{base_name}.pdf")

# ================= DATA =================
with tab3:
    st.subheader("📊 CSV ↔ Excel")

    uploaded_file = st.file_uploader("Upload File", type=["csv","xlsx"])
    option = st.selectbox("Conversion", ["CSV to Excel","Excel to CSV"])

    if uploaded_file:
        file_path, base_name = save_temp_file(uploaded_file)

        if st.button("🚀 Convert Data"):
            if option == "CSV to Excel":
                df = pd.read_csv(file_path)
                output = file_path + ".xlsx"
                df.to_excel(output,index=False)
            else:
                df = pd.read_excel(file_path)
                output = file_path + ".csv"
                df.to_csv(output,index=False)

            with open(output,"rb") as f:
                st.download_button(
                    "📥 Download",
                    f,
                    file_name=f"{base_name}.{output.split('.')[-1]}"
                )

# ================= VIDEO =================
with tab4:
    st.subheader("🎥 Video Converter")

    if not VIDEO_ENABLED:
        st.warning("⚠️ Video conversion disabled on cloud")
    else:
        uploaded_file = st.file_uploader("Upload Video", type=["mp4","avi","mov"])
        format_choice = st.selectbox("Convert to", ["mp4","avi","mov"])

        if uploaded_file:
            file_path, base_name = save_temp_file(uploaded_file)

            if st.button("🚀 Convert Video"):
                clip = VideoFileClip(file_path)
                output = file_path + f".{format_choice}"
                clip.write_videofile(output)

                with open(output,"rb") as f:
                    st.download_button("📥 Download", f, file_name=f"{base_name}.{format_choice}")

# ================= AUDIO =================
with tab5:
    st.subheader("🎵 Audio Converter")

    if not VIDEO_ENABLED:
        st.warning("⚠️ Audio conversion disabled on cloud")
    else:
        uploaded_file = st.file_uploader("Upload Audio", type=["mp3","wav"])
        format_choice = st.selectbox("Convert to", ["mp3","wav"])

        if uploaded_file:
            file_path, base_name = save_temp_file(uploaded_file)

            if st.button("🚀 Convert Audio"):
                audio = AudioFileClip(file_path)
                output = file_path + f".{format_choice}"
                audio.write_audiofile(output)

                with open(output,"rb") as f:
                    st.download_button("📥 Download", f, file_name=f"{base_name}.{format_choice}")
